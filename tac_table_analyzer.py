#!/usr/bin/env python3
"""
Specialized TAC Document Table Extractor
Focuses specifically on extracting contact tables from TAC documents pages 7-10

This version:
1. Extracts ALL tables from documents and preserves them
2. Specifically targets BFT and CFT contact information
3. Maps tables to their corresponding sections more accurately
4. Preserves table structure for contact information
"""

import json
import os
import re
from datetime import datetime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

def extract_all_tables_with_context(file_path):
    """Extract all tables from document with surrounding context"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(file_path)
        all_tables = []
        current_section = "unknown"
        
        # Get all paragraphs and tables in document order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find the paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            # Check if this looks like a section heading
                            if any(keyword in text.upper() for keyword in ["SECTION", "BUSINESS FUNCTION", "BFT", "CFT", "RECOVERY TEAM"]):
                                current_section = text
                        break
                        
            elif element.tag.endswith('tbl'):  # Table
                # Find the table object
                for table in doc.tables:
                    if table._element == element:
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                row_data.append(cell_text)
                            if any(cell.strip() for cell in row_data):  # Not empty row
                                table_data.append(row_data)
                        
                        if table_data:
                            all_tables.append({
                                "data": table_data,
                                "context": current_section,
                                "row_count": len(table_data),
                                "col_count": len(table_data[0]) if table_data else 0
                            })
                        break
        
        return {"tables": all_tables}
        
    except Exception as e:
        return {"error": f"Error extracting tables: {str(e)}"}

def analyze_table_content(table_info):
    """Analyze table content to determine if it contains contact information"""
    table_data = table_info["data"]
    context = table_info["context"]
    
    analysis = {
        "is_contact_table": False,
        "has_names": False,
        "has_phone_numbers": False,
        "has_emails": False,
        "has_titles": False,
        "likely_contact_type": None,
        "sample_data": table_data[:3] if len(table_data) > 0 else []
    }
    
    # Check context for BFT/CFT indicators
    context_lower = context.lower()
    if "bft" in context_lower or "business function" in context_lower:
        analysis["likely_contact_type"] = "BFT"
    elif "cft" in context_lower or "cross functional" in context_lower:
        analysis["likely_contact_type"] = "CFT"
    
    # Analyze table content
    all_text = " ".join([" ".join(row) for row in table_data]).lower()
    
    # Check for contact indicators
    if any(indicator in all_text for indicator in ["name", "title", "phone", "mobile", "email", "contact"]):
        analysis["has_contact_headers"] = True
    
    # Check for phone number patterns
    phone_pattern = r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b|\(\d{3}\)\s?\d{3}[-.\s]?\d{4}|\+\d+'
    if re.search(phone_pattern, all_text):
        analysis["has_phone_numbers"] = True
    
    # Check for email patterns
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    if re.search(email_pattern, all_text):
        analysis["has_emails"] = True
    
    # Check for probable names (capitalized words that look like names)
    for row in table_data:
        for cell in row:
            # Look for patterns like "First Last" or "First Middle Last"
            words = cell.strip().split()
            if len(words) >= 2 and len(words) <= 4:
                if all(word[0].isupper() and word[1:].islower() for word in words if len(word) > 1):
                    if not any(exclude in cell.lower() for exclude in ["section", "team", "department", "function", "title"]):
                        analysis["has_names"] = True
                        break
        if analysis["has_names"]:
            break
    
    # Check for titles/roles
    title_indicators = ["manager", "director", "lead", "coordinator", "specialist", "engineer", "analyst"]
    if any(title in all_text for title in title_indicators):
        analysis["has_titles"] = True
    
    # Determine if this is likely a contact table
    contact_score = 0
    if analysis.get("has_contact_headers", False): contact_score += 2
    if analysis["has_names"]: contact_score += 2
    if analysis["has_phone_numbers"]: contact_score += 2
    if analysis["has_emails"]: contact_score += 1
    if analysis["has_titles"]: contact_score += 1
    if analysis["likely_contact_type"]: contact_score += 1
    
    analysis["is_contact_table"] = contact_score >= 3
    analysis["contact_score"] = contact_score
    
    return analysis

def extract_structured_contacts_from_table(table_data):
    """Extract structured contact information from a validated contact table"""
    contacts = []
    
    if not table_data or len(table_data) < 2:
        return contacts
    
    # Try to identify header row
    headers = []
    header_row_idx = -1
    
    for i, row in enumerate(table_data[:3]):  # Check first 3 rows for headers
        row_text = " ".join(row).lower()
        if any(header in row_text for header in ["name", "title", "phone", "mobile", "email"]):
            headers = [col.strip().lower() for col in row]
            header_row_idx = i
            break
    
    # If no headers found, use positional mapping
    if header_row_idx == -1 and len(table_data[0]) >= 2:
        headers = ["name", "title"]
        if len(table_data[0]) >= 3: headers.append("phone")
        if len(table_data[0]) >= 4: headers.append("mobile")
        if len(table_data[0]) >= 5: headers.append("email")
        if len(table_data[0]) >= 6: headers.append("location")
        header_row_idx = -1  # Start from first row
    
    # Extract contact data
    start_row = header_row_idx + 1 if header_row_idx >= 0 else 0
    
    for row_idx, row in enumerate(table_data[start_row:]):
        if not any(cell.strip() for cell in row):  # Skip empty rows
            continue
        
        contact = {
            "type": "team_member",
            "name": "",
            "title": "",
            "phone": [],
            "mobile": [],
            "email": [],
            "location": "",
            "description": "",
            "raw_data": row,
            "row_index": start_row + row_idx
        }
        
        # Map data to fields
        for col_idx, cell in enumerate(row):
            cell = cell.strip()
            if not cell:
                continue
            
            # Get header for this column
            header = headers[col_idx] if col_idx < len(headers) else f"field_{col_idx}"
            
            # Extract phone and email patterns
            phone_pattern = r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b|\(\d{3}\)\s?\d{3}[-.\s]?\d{4}|\+\d+[-.\s]?\d+'
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            
            phones = re.findall(phone_pattern, cell)
            emails = re.findall(email_pattern, cell)
            
            # Map based on header or content
            if "name" in header or col_idx == 0:
                if not any(exclude in cell.lower() for exclude in ["section", "team", "department", "phone", "email", "@"]):
                    contact["name"] = cell
            elif "title" in header or "role" in header or col_idx == 1:
                contact["title"] = cell
            elif "phone" in header and "mobile" not in header:
                if phones:
                    contact["phone"].extend(phones)
                elif not emails:
                    contact["phone"].append(cell)
            elif "mobile" in header or "cell" in header:
                if phones:
                    contact["mobile"].extend(phones)
                elif not emails:
                    contact["mobile"].append(cell)
            elif "email" in header or emails:
                contact["email"].extend(emails)
                if not emails and "@" not in cell:  # Not an email pattern
                    contact["email"].append(cell)
            elif "location" in header or "site" in header:
                contact["location"] = cell
            elif "description" in header or "notes" in header:
                contact["description"] = cell
            else:
                # Auto-detect content type
                if emails:
                    contact["email"].extend(emails)
                elif phones:
                    contact["phone"].extend(phones)
                elif col_idx >= 2 and not contact["location"] and len(cell) > 5:
                    contact["location"] = cell
        
        # Only add if we have a reasonable name
        if contact["name"] and len(contact["name"]) > 2:
            contacts.append(contact)
    
    return contacts

def extract_tac_specific_data():
    """Extract data specifically focusing on TAC documents and table preservation"""
    
    file_mapping = {
        "TAC-Global BCP (3).docx": "tac",
        "CX Business Continuity Plan TAC.docx": "cx-tac"
    }
    
    results = {}
    
    for filename, plan_id in file_mapping.items():
        if os.path.exists(filename):
            print(f"\n🔍 Analyzing {filename} for table content...")
            
            # Extract all tables with context
            extraction_result = extract_all_tables_with_context(filename)
            
            if "error" in extraction_result:
                print(f"  ❌ Error: {extraction_result['error']}")
                continue
            
            tables = extraction_result["tables"]
            print(f"  📊 Found {len(tables)} tables total")
            
            # Analyze each table
            contact_tables = []
            all_contacts = []
            
            for i, table_info in enumerate(tables):
                analysis = analyze_table_content(table_info)
                print(f"  📋 Table {i+1}: {table_info['row_count']}x{table_info['col_count']} - Context: {table_info['context'][:50]}...")
                print(f"      Contact Score: {analysis['contact_score']}/7 - Is Contact Table: {analysis['is_contact_table']}")
                
                if analysis["is_contact_table"]:
                    print(f"      ✅ Identified as {analysis.get('likely_contact_type', 'contact')} table")
                    contact_tables.append({
                        "table_index": i,
                        "analysis": analysis,
                        "data": table_info["data"]
                    })
                    
                    # Extract contacts from this table
                    contacts = extract_structured_contacts_from_table(table_info["data"])
                    all_contacts.extend(contacts)
                    print(f"      👥 Extracted {len(contacts)} contacts")
                else:
                    print(f"      ⚠️  Not identified as contact table")
                
                # Show sample data for review
                print(f"      Sample: {table_info['data'][0] if table_info['data'] else 'Empty'}")
            
            results[plan_id] = {
                "filename": filename,
                "total_tables": len(tables),
                "contact_tables": contact_tables,
                "all_contacts": all_contacts,
                "contact_count": len(all_contacts)
            }
            
            print(f"  🎯 Final result: {len(all_contacts)} total contacts from {len(contact_tables)} contact tables")
        
        else:
            print(f"❌ File not found: {filename}")
    
    return results

if __name__ == "__main__":
    if not DOCX_AVAILABLE:
        print("❌ Error: python-docx is required but not installed.")
        print("Install it with: pip install python-docx")
        exit(1)
    
    print("🎯 TAC Document Table Analysis")
    print("📋 Focus: Extract and analyze ALL tables, especially contact tables from pages 7-10")
    print()
    
    results = extract_tac_specific_data()
    
    if results:
        print("\n" + "="*60)
        print("📊 ANALYSIS SUMMARY")
        print("="*60)
        
        for plan_id, data in results.items():
            print(f"\n📄 {data['filename']}:")
            print(f"  📊 Total Tables: {data['total_tables']}")
            print(f"  👥 Contact Tables: {len(data['contact_tables'])}")
            print(f"  🎯 Total Contacts: {data['contact_count']}")
            
            if data['all_contacts']:
                print(f"  📋 Sample Contacts:")
                for i, contact in enumerate(data['all_contacts'][:3]):
                    print(f"    {i+1}. {contact['name']} - {contact['title']}")
        
        # Save detailed results
        with open('tac_table_analysis.json', 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        
        print(f"\n✅ Detailed analysis saved to: tac_table_analysis.json")
    else:
        print("\n❌ No results to analyze")
