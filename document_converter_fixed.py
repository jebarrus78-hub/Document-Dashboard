#!/usr/bin/env python3
"""
Fixed BCP Document Converter
Properly identifies and extracts only actual team member contact information

This version fixes:
1. False positives where procedures are identified as team members
2. Better table header detection for contact tables
3. More strict validation of what constitutes a person vs. procedure
"""

import json
import os
import re
from datetime import datetime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

def is_person_name(text):
    """Check if text looks like a person's name"""
    if not text or len(text) < 2:
        return False
    
    # Common non-person indicators
    non_person_words = [
        "return", "normal", "operations", "activate", "section", "page", "step",
        "procedure", "process", "system", "tool", "application", "recovery",
        "business", "continuity", "plan", "when", "if", "the", "this", "that",
        "contact", "engage", "determine", "communicate", "situation", "team",
        "activation", "equipment", "outage", "building", "location", "personnel"
    ]
    
    text_lower = text.lower()
    if any(word in text_lower for word in non_person_words):
        return False
    
    # Should be reasonable length for a name
    if len(text) > 50:
        return False
    
    # Should contain mostly letters and spaces
    if not re.match(r'^[A-Za-z\s\.\-\']+$', text):
        return False
    
    # Should have at least one capital letter (names are usually capitalized)
    if not any(c.isupper() for c in text):
        return False
    
    # Should look like "First Last" or "First Middle Last" pattern
    words = text.split()
    if len(words) < 2 or len(words) > 4:
        return False
    
    # Each word should be a reasonable length for name parts
    for word in words:
        if len(word) < 2 or len(word) > 20:
            return False
    
    return True

def is_contact_table(table_data):
    """Determine if a table contains contact information"""
    if not table_data or len(table_data) < 2:
        return False
    
    # Check first few rows for contact-related headers
    header_indicators = []
    for i, row in enumerate(table_data[:3]):  # Check first 3 rows
        row_text = " ".join(row).lower()
        header_indicators.append(any(header in row_text for header in [
            "name", "title", "phone", "mobile", "email", "contact", "role", "team member"
        ]))
    
    # At least one of the first 3 rows should have contact headers
    if not any(header_indicators):
        return False
    
    # Check if the table has actual person names (not just procedures)
    person_count = 0
    for row in table_data[1:6]:  # Check next 5 rows after potential header
        if row and len(row) > 0:
            first_cell = row[0].strip()
            if is_person_name(first_cell):
                person_count += 1
    
    # Should have at least 2 people to be considered a contact table
    return person_count >= 2

def extract_contact_table(table_data):
    """Extract structured contact information from validated contact table"""
    if not is_contact_table(table_data):
        return []
    
    contacts = []
    headers = []
    header_row_index = -1
    
    # Find the header row
    for i, row in enumerate(table_data):
        row_text = " ".join(row).lower()
        if any(header in row_text for header in ["name", "title", "phone", "mobile", "email", "contact"]):
            headers = [col.lower().strip() for col in row]
            header_row_index = i
            break
    
    if header_row_index == -1:
        # Try to infer structure if no clear headers
        if len(table_data[0]) >= 3:
            headers = ["name", "title", "contact_info"]
            for i in range(3, len(table_data[0])):
                headers.append(f"field_{i}")
            header_row_index = -1  # Start from first row
    
    # Extract contact data
    start_row = header_row_index + 1 if header_row_index >= 0 else 0
    
    for row in table_data[start_row:]:
        if not any(cell.strip() for cell in row):  # Skip empty rows
            continue
        
        # First cell should be a person's name
        if not row or not is_person_name(row[0].strip()):
            continue
        
        contact = {
            "type": "team_member",
            "name": "",
            "title": "",
            "phone": [],
            "mobile": [],
            "email": [],
            "location": "",
            "description": "",
            "raw_data": row
        }
        
        # Map row data to contact fields
        for i, cell in enumerate(row):
            if i >= len(headers):
                break
                
            cell = cell.strip()
            if not cell:
                continue
                
            header = headers[i] if i < len(headers) else f"field_{i}"
            
            # Extract emails
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, cell)
            if emails:
                contact["email"].extend(emails)
            
            # Extract phone numbers
            phone_patterns = [
                r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
                r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
                r'\b1[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
                r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            ]
            
            phones = []
            for pattern in phone_patterns:
                phones.extend(re.findall(pattern, cell))
            
            # Map to appropriate fields based on header and content
            if i == 0 or "name" in header:
                contact["name"] = cell
            elif i == 1 or "title" in header or "role" in header:
                contact["title"] = cell
            elif "phone" in header and "mobile" not in header:
                if phones:
                    contact["phone"].extend(phones)
                elif not emails and len(cell) > 5:  # Likely phone number
                    contact["phone"].append(cell)
            elif "mobile" in header or "cell" in header:
                if phones:
                    contact["mobile"].extend(phones)
                elif not emails and len(cell) > 5:  # Likely mobile number
                    contact["mobile"].append(cell)
            elif "location" in header or "site" in header:
                contact["location"] = cell
            elif "description" in header or "notes" in header:
                contact["description"] = cell
            elif "email" in header:
                if not emails:  # If no email pattern found, treat as email anyway
                    contact["email"].append(cell)
            else:
                # For other fields, try to infer content type
                if emails:
                    contact["email"].extend(emails)
                elif phones:
                    if "mobile" in cell.lower() or len(contact["phone"]) > 0:
                        contact["mobile"].extend(phones)
                    else:
                        contact["phone"].extend(phones)
                elif i == len(row) - 1 and len(cell) > 10:  # Last column, longer text
                    if not contact["location"] and any(word in cell.lower() for word in ["office", "building", "floor", "room"]):
                        contact["location"] = cell
                    elif not contact["description"]:
                        contact["description"] = cell
        
        # Only add contacts with valid names
        if contact["name"] and is_person_name(contact["name"]):
            contacts.append(contact)
    
    return contacts

def extract_fixed_document_structure(file_path):
    """Extract document structure with fixed contact table parsing"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(file_path)
        structure = {
            "sections": {},
            "toc": [],
            "tables": []
        }
        
        current_section_key = None
        current_section_title = None
        current_content = []
        
        # First pass: extract paragraphs and identify sections
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # Check if this is a section heading
            is_heading = (
                ("heading" in style_name.lower() if style_name else False) or
                (style_name in ["Title", "Heading 1", "Heading 2"] if style_name else False) or
                any(keyword in text.upper() for keyword in ["SECTION", "APPENDIX"]) or
                (text.isupper() and len(text) > 10 and len(text) < 100)
            )
            
            if is_heading:
                # Save previous section
                if current_section_key and current_content:
                    structure["sections"][current_section_key] = {
                        "title": current_section_title,
                        "type": "general",
                        "content": current_content.copy()
                    }
                
                # Start new section
                current_section_key = text.lower().replace(" ", "_").replace(":", "").replace("(", "").replace(")", "").replace("-", "_").replace(",", "")
                current_section_title = text
                current_content = []
                
                # Determine section level
                level = 1 if "SECTION" in text.upper() or "APPENDIX" in text.upper() else 2
                
                structure["toc"].append({
                    "key": current_section_key,
                    "title": text,
                    "level": level
                })
                
            else:
                # Add content to current section
                if current_section_key:
                    current_content.append({
                        "type": "text",
                        "content": text
                    })
        
        # Save the last section
        if current_section_key and current_content:
            structure["sections"][current_section_key] = {
                "title": current_section_title,
                "type": "general", 
                "content": current_content
            }
        
        # Second pass: extract and process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            
            if table_data:
                structure["tables"].append(table_data)
                
                # Only process tables that are actually contact tables
                if is_contact_table(table_data):
                    contacts = extract_contact_table(table_data)
                    
                    if contacts:
                        # Look for BFT or CFT related sections
                        table_context = " ".join([" ".join(row) for row in table_data[:2]]).lower()
                        
                        target_section = None
                        if "bft" in table_context or "business function" in table_context:
                            target_section = "business_function_team_bft"
                        elif "cft" in table_context or "cross functional" in table_context:
                            target_section = "cross_functional_team_cft"
                        
                        # Also check current section context
                        if not target_section and current_section_key:
                            if "business_function" in current_section_key or "bft" in current_section_key:
                                target_section = current_section_key
                        
                        if target_section and target_section in structure["sections"]:
                            structure["sections"][target_section]["type"] = "contacts"
                            if not structure["sections"][target_section]["content"]:
                                structure["sections"][target_section]["content"] = []
                            structure["sections"][target_section]["content"].extend(contacts)
                        
                        elif contacts:  # Create section if needed
                            section_key = target_section or f"team_contacts_{len(structure['sections'])}"
                            structure["sections"][section_key] = {
                                "title": "Business Function Team (BFT)" if "bft" in table_context else "Team Contacts",
                                "type": "contacts",
                                "content": contacts
                            }
        
        return structure
        
    except Exception as e:
        return {"error": f"Error extracting structure: {str(e)}"}

def convert_with_fixed_extraction():
    """Convert documents with fixed contact extraction"""
    
    file_mapping = {
        "CE_Global_BCP.docx": "ce",
        "CMS Global BCP (2).docx": "cms", 
        "CX Business Continuity Plan TAC.docx": "cx-tac",
        "CX Labs (Global) BCP (2).docx": "cx-labs",
        "Global LSC Business Continuity Plan (BCP).docx": "lsc",
        "Proactive Services BCP (3).docx": "proactive",
        "Sourced Services (Global) BCP (3).docx": "sourced",
        "TAC-Global BCP (3).docx": "tac"
    }
    
    data = {
        "bcpPlans": {},
        "metadata": {
            "version": "5.0",
            "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
            "extractionMethod": "Fixed contact table extraction - person validation",
            "totalPlans": 0
        }
    }
    
    for filename, plan_id in file_mapping.items():
        if os.path.exists(filename):
            print(f"Processing {filename} with fixed extraction...")
            
            structure = extract_fixed_document_structure(filename)
            
            if "error" in structure:
                print(f"  Error: {structure['error']}")
                continue
            
            # Build the plan data
            plan_data = {
                "id": plan_id,
                "title": filename.replace(".docx", ""),
                "fileName": filename,
                "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
                "tableOfContents": structure["toc"],
                "sections": structure["sections"]
            }
            
            data["bcpPlans"][plan_id] = plan_data
            
            print(f"  ✓ Extracted {len(structure['toc'])} TOC items")
            print(f"  ✓ Processed {len(structure['sections'])} sections")
            print(f"  ✓ Found {len(structure['tables'])} total tables")
            
            # Count actual team member contacts
            contact_count = 0
            contact_tables = 0
            if isinstance(structure["sections"], dict):
                for section_key, section in structure["sections"].items():
                    if section["type"] == "contacts":
                        contact_tables += 1
                        team_members = [item for item in section["content"] if item.get("type") == "team_member"]
                        contact_count += len(team_members)
            
            if contact_count > 0:
                print(f"  ✓ Extracted {contact_count} verified team members from {contact_tables} contact tables")
            else:
                print(f"  ⚠ No team member contacts found")
        else:
            print(f"File not found: {filename}")
    
    data["metadata"]["totalPlans"] = len(data["bcpPlans"])
    
    # Save the fixed data
    output_file = "bcp-data-fixed.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    print(f"\n✅ Fixed extraction completed!")
    print(f"📄 Output saved to: {output_file}")
    print(f"📊 Total plans processed: {data['metadata']['totalPlans']}")
    
    return data

if __name__ == "__main__":
    if not DOCX_AVAILABLE:
        print("❌ Error: python-docx is required but not installed.")
        print("Install it with: pip install python-docx")
        exit(1)
    
    print("🔧 Starting fixed BCP document conversion...")
    print("📋 Focus: Proper person name validation and contact table detection")
    print()
    
    result = convert_with_fixed_extraction()
    
    if result:
        print("\n🎉 Fixed conversion completed successfully!")
        print("The fixed data includes:")
        print("  • Validated person names only")
        print("  • Proper contact table detection")
        print("  • No procedure text as team members")
        print("  • Clean structured contact information")
