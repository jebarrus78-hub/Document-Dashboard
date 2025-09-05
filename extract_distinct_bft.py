#!/usr/bin/env python3
"""
Extract distinct BFT/CFT contact information from each business line document
and create accurate BCP data with unique contact information per line of business.
"""

import json
import os
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    print("❌ python-docx not found. Installing...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.table import Table

class DistinctBFTExtractor:
    def __init__(self):
        self.docx_files = {
            "ce": "CE_Global_BCP.docx",
            "cms": "CMS Global BCP (2).docx", 
            "cx-labs": "CX Labs (Global) BCP (2).docx",
            "lsc": "Global LSC Business Continuity Plan (BCP).docx",
            "proactive": "Proactive Services BCP (3).docx",
            "sourced": "Sourced Services (Global) BCP (3).docx",
            "tac": "TAC-Global BCP (3).docx"
        }
        self.results = {}
        
    def clean_text(self, text):
        """Clean and normalize text content."""
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text.strip())
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text

    def extract_contact_from_row(self, row_data, headers):
        """Extract contact information from a table row."""
        contact = {
            "type": "team_member",
            "name": "",
            "userid": "",
            "raw_name": "",
            "title": "",
            "location": "",
            "phone": [],
            "mobile": [],
            "email": [],
            "description": ""
        }
        
        # Map headers to contact fields
        for i, header in enumerate(headers):
            if i >= len(row_data):
                continue
                
            header_lower = header.lower()
            value = self.clean_text(row_data[i])
            
            if not value or value.lower() in ['none', 'n/a', 'tbd', '']:
                continue
            
            # Name and User ID
            if 'name' in header_lower or 'user id' in header_lower:
                contact["raw_name"] = value
                # Extract name and userid
                if '(' in value and ')' in value:
                    name_part = value.split('(')[0].strip()
                    userid_match = re.search(r'\(([^)]+)\)', value)
                    contact["name"] = name_part
                    contact["userid"] = userid_match.group(1) if userid_match else ""
                else:
                    contact["name"] = value
                    
            # Title/Role
            elif 'title' in header_lower or 'role' in header_lower:
                contact["title"] = value
                
            # Location
            elif 'location' in header_lower:
                contact["location"] = value
                
            # Phone numbers
            elif 'phone' in header_lower and 'mobile' not in header_lower:
                if value and value not in ['None Listed', 'N/A']:
                    contact["phone"].append(value)
                    
            # Mobile numbers
            elif 'mobile' in header_lower:
                if value and value not in ['None Listed', 'N/A']:
                    contact["mobile"].append(value)
                    
            # Email
            elif 'email' in header_lower:
                if value and '@' in value:
                    contact["email"].append(value)
        
        # Only return contacts with at least a name
        if contact["name"] and len(contact["name"]) > 1:
            return contact
        return None

    def is_bft_cft_table(self, table, section_context=""):
        """Determine if a table contains BFT or CFT contact information."""
        if not table.rows or len(table.rows) < 2:
            return False
            
        # Check headers for contact-related content
        header_row = table.rows[0]
        header_text = ""
        for cell in header_row.cells:
            header_text += " " + self.clean_text(cell.text).lower()
        
        # Look for contact table indicators
        contact_indicators = ['name', 'user id', 'title', 'role', 'phone', 'mobile', 'location']
        has_contact_headers = any(indicator in header_text for indicator in contact_indicators)
        
        # Check for sufficient data rows
        data_rows = 0
        for row_idx, row in enumerate(table.rows[1:], 1):  # Skip header
            row_text = ""
            for cell in row.cells:
                row_text += self.clean_text(cell.text)
            if row_text.strip():
                data_rows += 1
        
        # Section context check
        section_lower = section_context.lower() if section_context else ""
        in_recovery_section = any(keyword in section_lower for keyword in ['recovery', 'team', 'bft', 'cft'])
        
        return has_contact_headers and data_rows >= 3 and (in_recovery_section or data_rows >= 10)

    def extract_bft_cft_from_document(self, line_id, file_path):
        """Extract BFT/CFT contact information from a single document."""
        print(f"\n📄 Extracting contacts from: {file_path} ({line_id})")
        
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        try:
            doc = Document(file_path)
            contacts = []
            current_section = ""
            table_count = 0
            
            # Process document elements using doc.paragraphs and doc.tables
            all_elements = []
            
            # Collect all paragraphs and tables in order
            for para in doc.paragraphs:
                all_elements.append(('paragraph', para))
            
            for table in doc.tables:
                all_elements.append(('table', table))
            
            # Process elements
            for element_type, element in all_elements:
                if element_type == 'paragraph':
                    para_text = self.clean_text(element.text)
                    if para_text and any(word in para_text.lower() for word in ['section', 'recovery', 'team', 'contact']):
                        current_section = para_text
                
                elif element_type == 'table':
                    table_count += 1
                    table = element
                    
                    if self.is_bft_cft_table(table, current_section):
                        print(f"  📋 Found contact table {table_count} in section: {current_section[:50]}...")
                        
                        # Extract headers
                        headers = []
                        if table.rows:
                            header_row = table.rows[0]
                            for cell in header_row.cells:
                                headers.append(self.clean_text(cell.text))
                        
                        # Extract contact data
                        table_contacts = []
                        for row_idx, row in enumerate(table.rows[1:], 1):
                            row_data = []
                            for cell in row.cells:
                                row_data.append(self.clean_text(cell.text))
                            
                            contact = self.extract_contact_from_row(row_data, headers)
                            if contact:
                                table_contacts.append(contact)
                        
                        print(f"    ✅ Extracted {len(table_contacts)} contacts")
                        contacts.extend(table_contacts)
            
            print(f"  🎯 Total contacts extracted: {len(contacts)}")
            
            return {
                'line_id': line_id,
                'filename': file_path,
                'total_contacts': len(contacts),
                'contacts': contacts,
                'extraction_date': datetime.now().isoformat()
            }
            
        except Exception as e:
            print(f"❌ Error processing {file_path}: {e}")
            return None

    def extract_all_distinct_contacts(self):
        """Extract distinct BFT/CFT contacts from all business line documents."""
        print("🔍 Extracting Distinct BFT/CFT Contacts from All Business Lines")
        print("=" * 70)
        
        for line_id, file_path in self.docx_files.items():
            result = self.extract_bft_cft_from_document(line_id, file_path)
            if result:
                self.results[line_id] = result
        
        return self.results

    def create_structured_bcp_data(self):
        """Create structured BCP data with distinct contacts per line of business."""
        print("\n📊 Creating Structured BCP Data...")
        
        # Load existing BCP structure
        bcp_data = {}
        try:
            with open('bcp-data-final.json', 'r', encoding='utf-8') as f:
                existing_data = json.load(f)
                bcp_data = existing_data.get('bcpPlans', {})
        except FileNotFoundError:
            print("❌ bcp-data-final.json not found. Creating new structure.")
        
        # Update each business line with its distinct contacts
        for line_id, contact_data in self.results.items():
            if line_id not in bcp_data:
                # Create basic structure if not exists
                bcp_data[line_id] = {
                    "title": contact_data['filename'].replace('.docx', ''),
                    "sections": {},
                    "tableOfContents": []
                }
            
            # Find or create BFT section
            bft_section_id = None
            for section_id, section in bcp_data[line_id].get('sections', {}).items():
                if 'business_function_team' in section_id.lower() or 'recovery_teams' in section_id.lower():
                    bft_section_id = section_id
                    break
            
            if not bft_section_id:
                bft_section_id = f"section_iii_recovery_teams_{line_id}"
                bcp_data[line_id]['sections'][bft_section_id] = {
                    "title": "SECTION III: Recovery Teams",
                    "content": []
                }
                
                # Add to TOC if not exists
                toc_exists = any('recovery' in item.get('title', '').lower() for item in bcp_data[line_id].get('tableOfContents', []))
                if not toc_exists:
                    bcp_data[line_id]['tableOfContents'].append({
                        "title": "SECTION III: Recovery Teams",
                        "key": bft_section_id,
                        "subsections": [
                            {
                                "title": "Business Function Team (BFT)",
                                "key": f"{bft_section_id}_bft"
                            }
                        ]
                    })
            
            # Update BFT section with distinct contacts
            bft_section = bcp_data[line_id]['sections'][bft_section_id]
            
            # Remove any existing generic BFT content
            bft_section['content'] = [item for item in bft_section.get('content', []) 
                                    if not (item.get('type') == 'contacts' and 'Business Function Team' in item.get('title', ''))]
            
            # Add the distinct BFT contacts
            if contact_data['contacts']:
                bft_section['content'].append({
                    "type": "subsection",
                    "title": "Business Function Team (BFT)",
                    "content": [
                        {
                            "type": "text",
                            "content": f"The following table lists the Business Function Team members for {bcp_data[line_id]['title']}. This team is responsible for coordinating business continuity activities and communications during an incident."
                        },
                        {
                            "type": "contacts",
                            "title": "Business Function Team Members",
                            "content": contact_data['contacts']
                        }
                    ]
                })
                
                print(f"  ✅ {line_id}: Added {len(contact_data['contacts'])} distinct BFT contacts")
            else:
                print(f"  ⚠️  {line_id}: No contacts found, keeping existing structure")
        
        return bcp_data

    def save_updated_bcp_data(self, bcp_data, output_file='bcp-data-distinct.json'):
        """Save the updated BCP data with distinct contacts."""
        try:
            final_data = {
                "metadata": {
                    "generated_date": datetime.now().isoformat(),
                    "version": "2.0_distinct_contacts",
                    "total_plans": len(bcp_data),
                    "extraction_method": "distinct_per_business_line"
                },
                "bcpPlans": bcp_data
            }
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(final_data, f, indent=2, ensure_ascii=False)
            
            print(f"✅ Updated BCP data saved to: {output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error saving BCP data: {e}")
            return False

    def generate_extraction_summary(self):
        """Generate a summary of the contact extraction."""
        print("\n" + "=" * 70)
        print("📊 DISTINCT CONTACT EXTRACTION SUMMARY")
        print("=" * 70)
        
        total_contacts = 0
        for line_id, data in self.results.items():
            contact_count = data['total_contacts']
            total_contacts += contact_count
            print(f"\n📋 {line_id.upper()}: {data['filename']}")
            print(f"  👥 Distinct Contacts: {contact_count}")
            
            if contact_count > 0:
                # Show sample contacts
                sample_contacts = data['contacts'][:3]
                for contact in sample_contacts:
                    name = contact['name']
                    title = contact['title']
                    phone = contact['phone'][0] if contact['phone'] else 'No phone'
                    print(f"    • {name} - {title} - {phone}")
                
                if contact_count > 3:
                    print(f"    ... and {contact_count - 3} more contacts")
        
        print(f"\n🎯 TOTAL DISTINCT CONTACTS EXTRACTED: {total_contacts}")
        print(f"📊 Average contacts per business line: {total_contacts / len(self.results):.1f}")

def main():
    """Main execution function."""
    extractor = DistinctBFTExtractor()
    
    # Extract distinct contacts from all documents
    results = extractor.extract_all_distinct_contacts()
    
    if results:
        # Generate summary
        extractor.generate_extraction_summary()
        
        # Create updated BCP data structure
        bcp_data = extractor.create_structured_bcp_data()
        
        # Save the results
        if extractor.save_updated_bcp_data(bcp_data):
            print(f"\n🎉 SUCCESS!")
            print("✅ Distinct BFT/CFT contacts extracted from each business line")
            print("✅ Updated BCP data structure created")
            print("✅ Each business line now has its own unique contact information")
            print("\n🔄 Next steps:")
            print("1. Review 'bcp-data-distinct.json' for accuracy")
            print("2. Replace 'bcp-data-final.json' with the new distinct data")
            print("3. Test the web application with the updated contact information")
            
            return True
    
    print("❌ No distinct contacts extracted.")
    return False

if __name__ == "__main__":
    main()
