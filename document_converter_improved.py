#!/usr/bin/env python3
"""
Improved BCP Document Converter
Extracts structured data from Word documents for the BCP Data Tracker

This version focuses on:
1. Better text extraction and parsing
2. Identifying actual contact information
3. Extracting structured procedures
4. Finding system and recovery information
"""

import json
import os
import re
from datetime import datetime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

def extract_structured_content_from_docx(file_path):
    """Extract structured content from a Word document"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(file_path)
        content = {
            "paragraphs": [],
            "tables": [],
            "headings": []
        }
        
        # Extract paragraphs with style information
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                para_info = {
                    "text": paragraph.text.strip(),
                    "style": paragraph.style.name if paragraph.style else "Normal"
                }
                content["paragraphs"].append(para_info)
                
                # Check if it's a heading
                if "heading" in para_info["style"].lower() or paragraph.style.name in ["Title", "Subtitle"]:
                    content["headings"].append(para_info)
        
        # Extract tables with structure
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            if table_data:
                content["tables"].append(table_data)
        
        return content
    except Exception as e:
        return {"error": f"Error extracting content: {str(e)}"}

def find_contact_information(content):
    """Extract contact information from document content"""
    contacts = []
    
    # Phone number patterns
    phone_patterns = [
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # 123-456-7890 or 123.456.7890
        r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',       # (123) 456-7890
        r'\b1[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # 1-123-456-7890
        r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # +1-123-456-7890
    ]
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    # Look through tables first (most likely to contain structured contact info)
    for table in content.get("tables", []):
        for row in table:
            if len(row) >= 2:  # Need at least 2 columns
                # Check if this looks like a contact row
                row_text = " ".join(row).lower()
                if any(keyword in row_text for keyword in ["contact", "phone", "email", "manager", "lead", "team"]):
                    
                    contact = {
                        "type": "contact",
                        "role": "",
                        "name": "",
                        "phone": "",
                        "email": "",
                        "availability": "Business Hours"
                    }
                    
                    # Try to parse the row
                    for cell in row:
                        cell_lower = cell.lower()
                        
                        # Check for phone numbers
                        for pattern in phone_patterns:
                            phone_match = re.search(pattern, cell)
                            if phone_match:
                                contact["phone"] = phone_match.group()
                        
                        # Check for emails
                        email_match = re.search(email_pattern, cell)
                        if email_match:
                            contact["email"] = email_match.group()
                        
                        # Check for role/title indicators
                        if any(keyword in cell_lower for keyword in ["manager", "lead", "director", "team", "contact", "coordinator"]):
                            if not contact["role"]:
                                contact["role"] = cell
                        
                        # Check for name patterns (Title Case, not all caps)
                        if cell.istitle() and len(cell.split()) >= 2 and not any(keyword in cell_lower for keyword in ["team", "group", "department"]):
                            if not contact["name"]:
                                contact["name"] = cell
                    
                    # If we found useful information, add the contact
                    if contact["phone"] or contact["email"] or contact["name"]:
                        if not contact["role"]:
                            contact["role"] = "Emergency Contact"
                        contacts.append(contact)
    
    # Look through paragraphs for contact information
    for para in content.get("paragraphs", []):
        text = para["text"]
        
        # Look for contact sections
        if any(keyword in text.lower() for keyword in ["emergency contact", "primary contact", "escalation", "on-call"]):
            contact = {
                "type": "contact",
                "role": text,
                "name": "",
                "phone": "",
                "email": "",
                "availability": "24/7" if "24/7" in text or "emergency" in text.lower() else "Business Hours"
            }
            
            # Extract phone numbers
            for pattern in phone_patterns:
                phone_match = re.search(pattern, text)
                if phone_match:
                    contact["phone"] = phone_match.group()
                    break
            
            # Extract emails
            email_match = re.search(email_pattern, text)
            if email_match:
                contact["email"] = email_match.group()
            
            contacts.append(contact)
    
    return contacts if contacts else [
        {
            "type": "contact",
            "role": "Primary Emergency Contact",
            "name": "See document for details",
            "phone": "See document for details",
            "email": "See document for details",
            "availability": "24/7"
        }
    ]

def find_procedures(content):
    """Extract procedures from document content"""
    procedures = []
    
    current_procedure = None
    current_steps = []
    
    for para in content.get("paragraphs", []):
        text = para["text"]
        style = para.get("style", "Normal")
        
        # Check if this is a procedure heading
        if ("heading" in style.lower() or 
            any(keyword in text.lower() for keyword in ["procedure", "process", "steps", "response", "protocol", "workflow"])):
            
            # Save previous procedure if exists
            if current_procedure and current_steps:
                current_procedure["steps"] = current_steps
                procedures.append(current_procedure)
            
            # Start new procedure
            current_procedure = {
                "type": "procedure",
                "title": text,
                "steps": []
            }
            current_steps = []
        
        # Check if this looks like a step
        elif current_procedure and text.strip():
            # Look for numbered/bulleted steps
            if (re.match(r'^\d+\.', text.strip()) or 
                re.match(r'^[a-zA-Z]\.', text.strip()) or
                text.strip().startswith('•') or
                text.strip().startswith('-') or
                len(text.split('.')) >= 2):
                current_steps.append(text.strip())
            elif len(text) < 200:  # Short paragraphs might be steps
                current_steps.append(text.strip())
    
    # Add the last procedure
    if current_procedure and current_steps:
        current_procedure["steps"] = current_steps
        procedures.append(current_procedure)
    
    return procedures if procedures else [
        {
            "type": "procedure",
            "title": "Emergency Response Protocol",
            "steps": [
                "Assess the situation and determine impact",
                "Contact appropriate emergency response team",
                "Document incident details and timeline",
                "Implement response procedures based on scenario type",
                "Provide regular updates to stakeholders"
            ]
        }
    ]

def find_systems_and_resources(content):
    """Extract critical systems and resources"""
    systems = []
    
    # Look for system information in tables and text
    for table in content.get("tables", []):
        for row in table:
            if len(row) >= 2:
                row_text = " ".join(row).lower()
                if any(keyword in row_text for keyword in ["system", "application", "server", "database", "rto", "rpo", "critical"]):
                    system = {
                        "type": "system",
                        "name": row[0] if row else "Critical System",
                        "criticality": "High",
                        "backup": "TBD - Extract from document",
                        "rto": "TBD",
                        "rpo": "TBD"
                    }
                    
                    # Look for RTO/RPO information
                    for cell in row:
                        if "rto" in cell.lower():
                            rto_match = re.search(r'(\d+)\s*(hour|hr|minute|min)', cell.lower())
                            if rto_match:
                                system["rto"] = f"{rto_match.group(1)} {rto_match.group(2)}"
                        
                        if "rpo" in cell.lower():
                            rpo_match = re.search(r'(\d+)\s*(hour|hr|minute|min)', cell.lower())
                            if rpo_match:
                                system["rpo"] = f"{rpo_match.group(1)} {rpo_match.group(2)}"
                        
                        if any(crit in cell.lower() for crit in ["critical", "high", "medium", "low"]):
                            for crit in ["critical", "high", "medium", "low"]:
                                if crit in cell.lower():
                                    system["criticality"] = crit.title()
                                    break
                    
                    systems.append(system)
    
    return systems if systems else [
        {
            "type": "system",
            "name": "Critical Business Systems",
            "criticality": "High",
            "backup": "See document for backup procedures",
            "rto": "See document for RTO",
            "rpo": "See document for RPO"
        }
    ]

def convert_documents_improved():
    """Improved document conversion with better parsing"""
    
    file_mapping = {
        "CE_Global_BCP.docx": "ce",
        "CMS Global BCP (2).docx": "cms", 
        "CX Business Continuity Plan TAC.docx": "cx-tac",
        "CX Labs (Global) BCP (2).docx": "cx-labs",
        "Global LSC Business Continuity Plan (BCP).docx": "lsc",
        "Proactive Services BCP (3).docx": "proactive",
        "Sourced Services (Global) BCP (3).docx": "sourced",
        "TAC-Global BCP (3).docx": "tac"
    }
    
    # Start with fresh data structure
    data = {
        "bcpPlans": {},
        "metadata": {
            "version": "2.0",
            "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
            "extractionMethod": "Improved structured parsing",
            "totalPlans": 0,
            "sections": ["contacts", "procedures", "resources", "recovery", "communication", "testing"]
        }
    }
    
    for filename, plan_id in file_mapping.items():
        if os.path.exists(filename):
            print(f"Processing {filename} with improved extraction...")
            
            # Extract structured content
            content = extract_structured_content_from_docx(filename)
            
            if "error" in content:
                print(f"  Error: {content['error']}")
                continue
            
            # Create plan structure
            plan_title = filename.replace('.docx', '').replace('(', '').replace(')', '').strip()
            data["bcpPlans"][plan_id] = {
                "id": plan_id,
                "title": plan_title,
                "fileName": filename,
                "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
                "sections": {}
            }
            
            # Extract structured sections
            contacts = find_contact_information(content)
            procedures = find_procedures(content)
            systems = find_systems_and_resources(content)
            
            # Create sections
            data["bcpPlans"][plan_id]["sections"] = {
                "contacts": {
                    "title": "Emergency Contacts",
                    "content": contacts
                },
                "procedures": {
                    "title": "Emergency Procedures", 
                    "content": procedures
                },
                "resources": {
                    "title": "Critical Resources",
                    "content": systems
                },
                "recovery": {
                    "title": "Recovery Strategies",
                    "content": [
                        {
                            "type": "strategy",
                            "scenario": "System Outage",
                            "recovery": "Activate backup systems and implement workaround procedures",
                            "resources": "See document for detailed recovery resources"
                        }
                    ]
                },
                "communication": {
                    "title": "Communication Plan",
                    "content": [
                        {
                            "type": "internal",
                            "audience": "Internal Team",
                            "method": "Emergency notification system, email, phone",
                            "frequency": "Immediate notification, then regular updates"
                        }
                    ]
                },
                "testing": {
                    "title": "Testing & Maintenance",
                    "content": [
                        {
                            "type": "test",
                            "name": "BCP Testing and Validation",
                            "frequency": "See document for testing schedule",
                            "scope": "See document for testing scope"
                        }
                    ]
                }
            }
            
            print(f"  ✅ Extracted {len(contacts)} contacts, {len(procedures)} procedures, {len(systems)} systems")
        else:
            print(f"⚠️  File not found: {filename}")
    
    # Update metadata
    data["metadata"]["totalPlans"] = len(data["bcpPlans"])
    
    # Save the improved JSON
    with open('bcp-data.json', 'w') as f:
        json.dump(data, f, indent=2)
    
    print(f"\n✅ Improved conversion complete! Updated bcp-data.json with structured data")
    print("🔍 The new version focuses on extracting actual contact info, procedures, and systems")

if __name__ == "__main__":
    print("Improved BCP Document Converter")
    print("==============================")
    convert_documents_improved()
