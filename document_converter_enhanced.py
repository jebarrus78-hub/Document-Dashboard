#!/usr/bin/env python3
"""
Enhanced BCP Document Converter
Specifically improved to extract BFT and CFT team contact tables with detailed information

This version focuses on:
1. Better table extraction for contact information
2. Specific handling of BFT (Business Function Team) and CFT (Cross-Functional Team) data
3. Structured extraction of names, titles, phone numbers, locations, etc.
"""

import json
import os
import re
from datetime import datetime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

def is_contact_table_header(row):
    """Check if a table row contains contact information headers"""
    row_text = " ".join(row).lower()
    contact_headers = [
        "name", "title", "phone", "mobile", "email", "location", 
        "role", "description", "contact", "team member", "function"
    ]
    return any(header in row_text for header in contact_headers)

def extract_contact_table(table_data):
    """Extract structured contact information from table data"""
    if not table_data or len(table_data) < 2:
        return []
    
    contacts = []
    headers = []
    header_row_index = -1
    
    # Find the header row
    for i, row in enumerate(table_data):
        if is_contact_table_header(row):
            headers = [col.lower().strip() for col in row]
            header_row_index = i
            break
    
    if header_row_index == -1:
        # No clear headers found, try to infer structure
        first_row = table_data[0]
        if len(first_row) >= 3:  # Assume name, title, contact info
            headers = ["name", "title", "contact_info"]
            for i in range(3, len(first_row)):
                headers.append(f"additional_{i}")
            header_row_index = -1  # Start from first row
    
    # Extract contact data
    start_row = header_row_index + 1 if header_row_index >= 0 else 0
    
    for row in table_data[start_row:]:
        if not any(cell.strip() for cell in row):  # Skip empty rows
            continue
            
        contact = {
            "type": "team_member",
            "name": "",
            "title": "",
            "phone": [],
            "mobile": [],
            "email": [],
            "location": "",
            "description": "",
            "raw_data": row
        }
        
        # Map row data to contact fields
        for i, cell in enumerate(row):
            if i >= len(headers):
                break
                
            cell = cell.strip()
            if not cell:
                continue
                
            header = headers[i]
            
            # Extract emails
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, cell)
            if emails:
                contact["email"].extend(emails)
            
            # Extract phone numbers
            phone_patterns = [
                r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
                r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
                r'\b1[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
                r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            ]
            
            phones = []
            for pattern in phone_patterns:
                phones.extend(re.findall(pattern, cell))
            
            # Map to appropriate fields based on header
            if "name" in header:
                contact["name"] = cell
            elif "title" in header or "role" in header:
                contact["title"] = cell
            elif "phone" in header and "mobile" not in header:
                if phones:
                    contact["phone"].extend(phones)
                else:
                    contact["phone"].append(cell)
            elif "mobile" in header or "cell" in header:
                if phones:
                    contact["mobile"].extend(phones)
                else:
                    contact["mobile"].append(cell)
            elif "location" in header:
                contact["location"] = cell
            elif "description" in header:
                contact["description"] = cell
            elif "email" in header:
                if not emails:  # If no email pattern found, treat as email anyway
                    contact["email"].append(cell)
            else:
                # For unclear headers, try to infer content type
                if emails:
                    contact["email"].extend(emails)
                elif phones:
                    contact["phone"].extend(phones)
                elif not contact["name"] and len(cell) > 2 and not any(char.isdigit() for char in cell):
                    contact["name"] = cell
                elif not contact["title"] and contact["name"]:
                    contact["title"] = cell
        
        # Only add contacts with at least a name
        if contact["name"]:
            contacts.append(contact)
    
    return contacts

def extract_enhanced_document_structure(file_path):
    """Extract document structure with enhanced contact table parsing"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(file_path)
        structure = {
            "sections": {},
            "toc": [],
            "tables": []
        }
        
        current_section_key = None
        current_section_title = None
        current_content = []
        
        # First pass: extract paragraphs and identify sections
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # Check if this is a section heading
            is_heading = (
                "heading" in style_name.lower() if style_name else False or
                style_name in ["Title", "Heading 1", "Heading 2"] if style_name else False or
                any(keyword in text.upper() for keyword in ["SECTION", "APPENDIX"]) or
                (text.isupper() and len(text) > 10 and len(text) < 100)
            )
            
            if is_heading:
                # Save previous section
                if current_section_key and current_content:
                    structure["sections"][current_section_key] = {
                        "title": current_section_title,
                        "type": "general",
                        "content": current_content.copy()
                    }
                
                # Start new section
                current_section_key = text.lower().replace(" ", "_").replace(":", "").replace("(", "").replace(")", "").replace("-", "_").replace(",", "")
                current_section_title = text
                current_content = []
                
                # Determine section level
                level = 1 if "SECTION" in text.upper() or "APPENDIX" in text.upper() else 2
                
                structure["toc"].append({
                    "key": current_section_key,
                    "title": text,
                    "level": level
                })
                
            else:
                # Add content to current section
                if current_section_key:
                    current_content.append({
                        "type": "text",
                        "content": text
                    })
        
        # Save the last section
        if current_section_key and current_content:
            structure["sections"][current_section_key] = {
                "title": current_section_title,
                "type": "general", 
                "content": current_content
            }
        
        # Second pass: extract and process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            
            if table_data:
                structure["tables"].append(table_data)
                
                # Try to determine which section this table belongs to
                table_context = " ".join([" ".join(row) for row in table_data[:2]]).lower()
                
                # Look for BFT or CFT related tables
                if any(keyword in table_context for keyword in ["bft", "business function team", "cft", "cross functional"]):
                    contacts = extract_contact_table(table_data)
                    
                    # Find the appropriate section
                    target_section = None
                    if "bft" in table_context or "business function" in table_context:
                        target_section = "business_function_team_bft"
                    elif "cft" in table_context or "cross functional" in table_context:
                        target_section = "cross_functional_team_cft"
                    
                    if target_section and target_section in structure["sections"]:
                        if not structure["sections"][target_section]["content"]:
                            structure["sections"][target_section]["content"] = []
                        
                        structure["sections"][target_section]["type"] = "contacts"
                        structure["sections"][target_section]["content"].extend(contacts)
                    
                    elif contacts:  # Create section if it doesn't exist
                        section_key = target_section or f"team_contacts_{len(structure['sections'])}"
                        structure["sections"][section_key] = {
                            "title": "Business Function Team (BFT)" if "bft" in table_context else "Team Contacts",
                            "type": "contacts",
                            "content": contacts
                        }
        
        return structure
        
    except Exception as e:
        return {"error": f"Error extracting structure: {str(e)}"}

def convert_with_enhanced_extraction():
    """Convert documents with enhanced contact extraction"""
    
    file_mapping = {
        "CE_Global_BCP.docx": "ce",
        "CMS Global BCP (2).docx": "cms", 
        "CX Business Continuity Plan TAC.docx": "cx-tac",
        "CX Labs (Global) BCP (2).docx": "cx-labs",
        "Global LSC Business Continuity Plan (BCP).docx": "lsc",
        "Proactive Services BCP (3).docx": "proactive",
        "Sourced Services (Global) BCP (3).docx": "sourced",
        "TAC-Global BCP (3).docx": "tac"
    }
    
    data = {
        "bcpPlans": {},
        "metadata": {
            "version": "4.0",
            "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
            "extractionMethod": "Enhanced contact table extraction",
            "totalPlans": 0
        }
    }
    
    for filename, plan_id in file_mapping.items():
        if os.path.exists(filename):
            print(f"Processing {filename} with enhanced extraction...")
            
            structure = extract_enhanced_document_structure(filename)
            
            if "error" in structure:
                print(f"  Error: {structure['error']}")
                continue
            
            # Build the plan data
            plan_data = {
                "id": plan_id,
                "title": filename.replace(".docx", ""),
                "fileName": filename,
                "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
                "tableOfContents": structure["toc"],
                "sections": structure["sections"]
            }
            
            data["bcpPlans"][plan_id] = plan_data
            
            print(f"  ✓ Extracted {len(structure['toc'])} TOC items")
            print(f"  ✓ Processed {len(structure['sections'])} sections")
            print(f"  ✓ Found {len(structure['tables'])} tables")
            
            # Count contacts in BFT/CFT sections
            contact_count = 0
            if isinstance(structure["sections"], dict):
                for section_key, section in structure["sections"].items():
                    if section["type"] == "contacts":
                        contact_count += len([item for item in section["content"] if item.get("type") == "team_member"])
            
            if contact_count > 0:
                print(f"  ✓ Extracted {contact_count} team member contacts")
        else:
            print(f"File not found: {filename}")
    
    data["metadata"]["totalPlans"] = len(data["bcpPlans"])
    
    # Save the enhanced data
    output_file = "bcp-data-enhanced.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    print(f"\n✅ Enhanced extraction completed!")
    print(f"📄 Output saved to: {output_file}")
    print(f"📊 Total plans processed: {data['metadata']['totalPlans']}")
    
    return data

if __name__ == "__main__":
    if not DOCX_AVAILABLE:
        print("❌ Error: python-docx is required but not installed.")
        print("Install it with: pip install python-docx")
        exit(1)
    
    print("🚀 Starting enhanced BCP document conversion...")
    print("📋 Focus: Improved BFT/CFT contact table extraction")
    print()
    
    result = convert_with_enhanced_extraction()
    
    if result:
        print("\n🎉 Conversion completed successfully!")
        print("The enhanced data includes:")
        print("  • Improved contact table parsing")
        print("  • Structured team member information")
        print("  • Names, titles, phone numbers, emails")
        print("  • Location and description data")
