#!/usr/bin/env python3
"""
BCP Document Converter
Converts Word documents to JSON format for the BCP Data Tracker

Usage:
    python3 document_converter.py

This script will:
1. Read all .docx files in the current directory
2. Extract text content from each document
3. Attempt to parse sections (contacts, procedures, etc.)
4. Update the bcp-data.json file with extracted content

Requirements:
    pip install python-docx
"""

import json
import os
from datetime import datetime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

def extract_text_from_docx(file_path):
    """Extract all text from a Word document"""
    if not DOCX_AVAILABLE:
        return "Could not extract content - python-docx not installed"
    
    try:
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return text_content
    except Exception as e:
        return f"Error extracting content: {str(e)}"

def parse_contacts_section(text_lines):
    """Attempt to parse contact information from text"""
    contacts = []
    
    # Look for phone numbers and emails
    for i, line in enumerate(text_lines):
        if any(keyword in line.lower() for keyword in ['contact', 'emergency', 'phone', 'email']):
            # Try to extract contact info from this and nearby lines
            contact_info = {
                "type": "contact",
                "role": line,
                "name": "TBD - Please update",
                "phone": "TBD - Please update", 
                "email": "TBD - Please update",
                "availability": "TBD - Please update"
            }
            
            # Look for phone numbers in nearby lines
            for j in range(max(0, i-2), min(len(text_lines), i+3)):
                line_text = text_lines[j]
                if any(char.isdigit() for char in line_text) and len(line_text.replace('-', '').replace('(', '').replace(')', '').replace(' ', '')) >= 10:
                    contact_info["phone"] = line_text
                if '@' in line_text:
                    contact_info["email"] = line_text
            
            contacts.append(contact_info)
    
    if not contacts:
        contacts.append({
            "type": "contact",
            "role": "Primary Emergency Contact",
            "name": "TBD - Please update from document",
            "phone": "TBD - Please update from document",
            "email": "TBD - Please update from document", 
            "availability": "24/7"
        })
    
    return contacts

def parse_procedures_section(text_lines):
    """Attempt to parse procedures from text"""
    procedures = []
    
    current_procedure = None
    steps = []
    
    for line in text_lines:
        # Look for procedure titles
        if any(keyword in line.lower() for keyword in ['procedure', 'process', 'step', 'response']):
            if current_procedure and steps:
                current_procedure["steps"] = steps
                procedures.append(current_procedure)
            
            current_procedure = {
                "type": "procedure",
                "title": line,
                "steps": []
            }
            steps = []
        elif line.strip() and current_procedure:
            steps.append(line)
    
    # Add the last procedure
    if current_procedure and steps:
        current_procedure["steps"] = steps
        procedures.append(current_procedure)
    
    if not procedures:
        procedures.append({
            "type": "procedure",
            "title": "Emergency Response Protocol",
            "steps": [
                "Assess the situation",
                "Contact emergency response team", 
                "Document incident details",
                "Implement appropriate response procedures"
            ]
        })
    
    return procedures

def create_section_content(text_lines, section_type):
    """Create content for a specific section based on extracted text"""
    if section_type == "contacts":
        return parse_contacts_section(text_lines)
    elif section_type == "procedures":
        return parse_procedures_section(text_lines)
    elif section_type == "resources":
        return [{
            "type": "system",
            "name": "Critical System (extracted from document)",
            "criticality": "High",
            "backup": "TBD - Please update from document",
            "rto": "TBD",
            "rpo": "TBD"
        }]
    elif section_type == "recovery":
        return [{
            "type": "strategy",
            "scenario": "System Failure",
            "recovery": "TBD - Please extract from document",
            "resources": "TBD - Please extract from document"
        }]
    elif section_type == "communication":
        return [{
            "type": "internal",
            "audience": "Internal Team",
            "method": "TBD - Please extract from document",
            "frequency": "TBD - Please extract from document"
        }]
    elif section_type == "testing":
        return [{
            "type": "test",
            "name": "BCP Testing",
            "frequency": "TBD - Please extract from document",
            "scope": "TBD - Please extract from document"
        }]
    else:
        return [{
            "type": "general",
            "title": "Extracted Content",
            "content": "\n".join(text_lines[:10])  # First 10 lines
        }]

def convert_documents():
    """Convert all Word documents in current directory to JSON"""
    
    # File mapping
    file_mapping = {
        "CE_Global_BCP.docx": "ce",
        "CMS Global BCP (2).docx": "cms", 
        "CX Business Continuity Plan TAC.docx": "cx-tac",
        "CX Labs (Global) BCP (2).docx": "cx-labs",
        "Global LSC Business Continuity Plan (BCP).docx": "lsc",
        "Proactive Services BCP (3).docx": "proactive",
        "Sourced Services (Global) BCP (3).docx": "sourced",
        "TAC-Global BCP (3).docx": "tac"
    }
    
    # Load existing JSON or create new structure
    try:
        with open('bcp-data.json', 'r') as f:
            data = json.load(f)
    except FileNotFoundError:
        data = {"bcpPlans": {}, "metadata": {}}
    
    sections = ["contacts", "procedures", "resources", "recovery", "communication", "testing"]
    
    for filename, plan_id in file_mapping.items():
        if os.path.exists(filename):
            print(f"Processing {filename}...")
            
            # Extract text from document
            text_lines = extract_text_from_docx(filename)
            
            if isinstance(text_lines, str):  # Error message
                print(f"  Error: {text_lines}")
                continue
            
            # Create plan structure
            plan_title = filename.replace('.docx', '').replace('(', '').replace(')', '').strip()
            if plan_id not in data["bcpPlans"]:
                data["bcpPlans"][plan_id] = {
                    "id": plan_id,
                    "title": plan_title,
                    "fileName": filename,
                    "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
                    "sections": {}
                }
            
            # Create sections
            for section in sections:
                section_content = create_section_content(text_lines, section)
                data["bcpPlans"][plan_id]["sections"][section] = {
                    "title": section.replace('_', ' ').title(),
                    "content": section_content
                }
            
            print(f"  ✅ Processed {len(text_lines)} lines of text")
        else:
            print(f"⚠️  File not found: {filename}")
    
    # Update metadata
    data["metadata"] = {
        "version": "1.1",
        "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
        "totalPlans": len(data["bcpPlans"]),
        "sections": sections
    }
    
    # Save updated JSON
    with open('bcp-data.json', 'w') as f:
        json.dump(data, f, indent=2)
    
    print(f"\n✅ Conversion complete! Updated bcp-data.json with {len(data['bcpPlans'])} plans")
    print("📝 Please review and update the extracted content as needed")

if __name__ == "__main__":
    print("BCP Document Converter")
    print("=====================")
    convert_documents()
