#!/usr/bin/env python3
"""
BCP Document Converter - Table of Contents Based
Extracts sections based on actual document structure and table of contents

This version:
1. Identifies section headings from the documents
2. Creates navigation based on actual BCP structure
3. Organizes content by document sections rather than generic categories
"""

import json
import os
import re
from datetime import datetime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

def extract_document_structure(file_path):
    """Extract document structure including headings and table of contents"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(file_path)
        structure = {
            "sections": {},
            "toc": [],
            "current_section": None,
            "current_content": []
        }
        
        current_section_key = None
        current_section_title = None
        current_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # Check if this is a major section heading
            if (any(keyword in text.upper() for keyword in ["SECTION", "APPENDIX"]) or
                "heading" in style_name.lower() or
                style_name in ["Title", "Heading 1", "Heading 2"] or
                text.isupper() and len(text) > 10 and len(text) < 100):
                
                # Save previous section if it exists
                if current_section_key and current_content:
                    structure["sections"][current_section_key] = {
                        "title": current_section_title,
                        "content": current_content.copy()
                    }
                
                # Start new section
                current_section_key = text.lower().replace(" ", "_").replace(":", "").replace("(", "").replace(")", "").replace("-", "_")
                current_section_title = text
                current_content = []
                
                # Add to table of contents
                structure["toc"].append({
                    "key": current_section_key,
                    "title": text,
                    "level": 1 if "SECTION" in text.upper() else 2
                })
                
            else:
                # Add content to current section
                if current_section_key:
                    current_content.append({
                        "type": "paragraph",
                        "text": text,
                        "style": style_name
                    })
        
        # Add the last section
        if current_section_key and current_content:
            structure["sections"][current_section_key] = {
                "title": current_section_title,
                "content": current_content
            }
        
        # Also extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    table_data.append(row_data)
            
            if table_data and current_section_key:
                if current_section_key not in structure["sections"]:
                    structure["sections"][current_section_key] = {
                        "title": "Table Data",
                        "content": []
                    }
                structure["sections"][current_section_key]["content"].append({
                    "type": "table",
                    "data": table_data
                })
        
        return structure
    except Exception as e:
        return {"error": f"Error extracting structure: {str(e)}"}

def extract_contacts_from_section(content):
    """Extract contact information from section content"""
    contacts = []
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    # Phone patterns
    phone_patterns = [
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
        r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
        r'\b1[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
        r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
    ]
    
    for item in content:
        if item["type"] == "paragraph":
            text = item["text"]
            
            # Look for contact information
            emails = re.findall(email_pattern, text)
            phones = []
            for pattern in phone_patterns:
                phones.extend(re.findall(pattern, text))
            
            if emails or phones or any(keyword in text.lower() for keyword in ["contact", "team", "manager", "lead"]):
                contact = {
                    "type": "contact",
                    "text": text,
                    "emails": emails,
                    "phones": phones
                }
                contacts.append(contact)
        
        elif item["type"] == "table":
            # Look for contact tables
            for row in item["data"]:
                row_text = " ".join(row)
                emails = re.findall(email_pattern, row_text)
                phones = []
                for pattern in phone_patterns:
                    phones.extend(re.findall(pattern, row_text))
                
                if emails or phones:
                    contact = {
                        "type": "contact",
                        "text": row_text,
                        "emails": emails,
                        "phones": phones
                    }
                    contacts.append(contact)
    
    return contacts

def convert_to_toc_based_structure():
    """Convert documents using table of contents structure"""
    
    file_mapping = {
        "CE_Global_BCP.docx": "ce",
        "CMS Global BCP (2).docx": "cms", 
        "CX Business Continuity Plan TAC.docx": "cx-tac",
        "CX Labs (Global) BCP (2).docx": "cx-labs",
        "Global LSC Business Continuity Plan (BCP).docx": "lsc",
        "Proactive Services BCP (3).docx": "proactive",
        "Sourced Services (Global) BCP (3).docx": "sourced",
        "TAC-Global BCP (3).docx": "tac"
    }
    
    data = {
        "bcpPlans": {},
        "metadata": {
            "version": "3.0",
            "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
            "extractionMethod": "Table of Contents based structure",
            "totalPlans": 0
        }
    }
    
    for filename, plan_id in file_mapping.items():
        if os.path.exists(filename):
            print(f"Extracting structure from {filename}...")
            
            structure = extract_document_structure(filename)
            
            if "error" in structure:
                print(f"  Error: {structure['error']}")
                continue
            
            plan_title = filename.replace('.docx', '').replace('(', '').replace(')', '').strip()
            
            # Convert sections to the format expected by the web app
            sections = {}
            for section_key, section_data in structure["sections"].items():
                
                # Extract different types of content based on section purpose
                if any(keyword in section_key for keyword in ["team", "contact", "communication"]):
                    # This is likely a contacts section
                    contacts = extract_contacts_from_section(section_data["content"])
                    sections[section_key] = {
                        "title": section_data["title"],
                        "type": "contacts",
                        "content": contacts if contacts else [{"type": "text", "content": f"See {section_data['title']} section in document"}]
                    }
                
                elif any(keyword in section_key for keyword in ["procedure", "process", "step", "response"]):
                    # This is likely a procedures section
                    procedures = []
                    current_procedure = {"type": "procedure", "title": section_data["title"], "steps": []}
                    
                    for item in section_data["content"]:
                        if item["type"] == "paragraph":
                            current_procedure["steps"].append(item["text"])
                    
                    procedures.append(current_procedure)
                    sections[section_key] = {
                        "title": section_data["title"],
                        "type": "procedures",
                        "content": procedures
                    }
                
                else:
                    # General content section
                    content_items = []
                    for item in section_data["content"]:
                        if item["type"] == "paragraph":
                            content_items.append({
                                "type": "text",
                                "content": item["text"]
                            })
                        elif item["type"] == "table":
                            content_items.append({
                                "type": "table",
                                "data": item["data"]
                            })
                    
                    sections[section_key] = {
                        "title": section_data["title"],
                        "type": "general",
                        "content": content_items
                    }
            
            data["bcpPlans"][plan_id] = {
                "id": plan_id,
                "title": plan_title,
                "fileName": filename,
                "lastUpdated": datetime.now().strftime("%Y-%m-%d"),
                "tableOfContents": structure["toc"],
                "sections": sections
            }
            
            print(f"  ✅ Extracted {len(structure['toc'])} TOC items and {len(sections)} sections")
        else:
            print(f"⚠️  File not found: {filename}")
    
    data["metadata"]["totalPlans"] = len(data["bcpPlans"])
    
    # Save the structured JSON
    with open('bcp-data-toc.json', 'w') as f:
        json.dump(data, f, indent=2)
    
    print(f"\n✅ TOC-based conversion complete! Created bcp-data-toc.json")
    print("📋 Structure now matches actual document organization")

if __name__ == "__main__":
    print("BCP Document Converter - Table of Contents Based")
    print("===============================================")
    convert_to_toc_based_structure()
