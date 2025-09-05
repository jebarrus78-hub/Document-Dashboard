#!/usr/bin/env python3
"""
BCP Document Converter
Converts Word documents to JSON format for the BCP Data Tracker webapp
"""

import os
import json
import re
from pathlib import Path
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document

class BCPDocumentConverter:
    def __init__(self, input_dir, output_dir="data"):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Common section keywords to identify different types of information
        self.section_keywords = {
            'contacts': [
                'contact', 'phone', 'email', 'personnel', 'emergency contact',
                'key personnel', 'escalation', 'notification', 'directory'
            ],
            'procedures': [
                'procedure', 'process', 'step', 'action', 'response', 'protocol',
                'emergency procedure', 'incident response', 'escalation procedure'
            ],
            'resources': [
                'resource', 'system', 'application', 'infrastructure', 'dependency',
                'critical system', 'technology', 'equipment', 'facility'
            ],
            'recovery': [
                'recovery', 'restore', 'backup', 'continuity', 'rto', 'rpo',
                'recovery time', 'recovery point', 'restoration', 'failover'
            ],
            'communication': [
                'communication', 'notification', 'announcement', 'stakeholder',
                'internal communication', 'external communication', 'media'
            ],
            'testing': [
                'test', 'drill', 'exercise', 'validation', 'maintenance',
                'review', 'update', 'assessment', 'audit'
            ]
        }

    def extract_text_from_docx(self, file_path):
        """Extract text content from Word document"""
        try:
            doc = Document(file_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'headers': []
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    style = para.style.name if para.style else 'Normal'
                    content['paragraphs'].append({
                        'text': para.text.strip(),
                        'style': style,
                        'is_heading': 'Heading' in style
                    })
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                content['tables'].append(table_data)
            
            return content
            
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return None

    def categorize_content(self, content):
        """Categorize content into sections based on keywords"""
        categorized = {
            'contacts': [],
            'procedures': [],
            'resources': [],
            'recovery': [],
            'communication': [],
            'testing': [],
            'general': []
        }
        
        # Process paragraphs
        for para in content['paragraphs']:
            text_lower = para['text'].lower()
            categorized_flag = False
            
            # Check against each category
            for category, keywords in self.section_keywords.items():
                if any(keyword in text_lower for keyword in keywords):
                    categorized[category].append({
                        'type': 'paragraph',
                        'content': para['text'],
                        'style': para['style'],
                        'is_heading': para['is_heading']
                    })
                    categorized_flag = True
                    break
            
            # If not categorized, add to general
            if not categorized_flag:
                categorized['general'].append({
                    'type': 'paragraph',
                    'content': para['text'],
                    'style': para['style'],
                    'is_heading': para['is_heading']
                })
        
        # Process tables - often contain contact info or procedures
        for table in content['tables']:
            if self.is_contact_table(table):
                categorized['contacts'].append({
                    'type': 'table',
                    'content': table,
                    'formatted': self.format_contact_table(table)
                })
            elif self.is_procedure_table(table):
                categorized['procedures'].append({
                    'type': 'table',
                    'content': table,
                    'formatted': self.format_procedure_table(table)
                })
            else:
                categorized['general'].append({
                    'type': 'table',
                    'content': table
                })
        
        return categorized

    def is_contact_table(self, table):
        """Determine if table contains contact information"""
        if not table or len(table) < 2:
            return False
        
        header_row = table[0]
        contact_indicators = ['name', 'phone', 'email', 'role', 'title', 'contact']
        
        return any(any(indicator in cell.lower() for indicator in contact_indicators) 
                  for cell in header_row)

    def is_procedure_table(self, table):
        """Determine if table contains procedure information"""
        if not table or len(table) < 2:
            return False
        
        header_row = table[0]
        procedure_indicators = ['step', 'action', 'procedure', 'process', 'task']
        
        return any(any(indicator in cell.lower() for indicator in procedure_indicators) 
                  for cell in header_row)

    def format_contact_table(self, table):
        """Format contact table for display"""
        if not table:
            return []
        
        headers = table[0]
        contacts = []
        
        for row in table[1:]:
            if any(cell.strip() for cell in row):  # Skip empty rows
                contact = {}
                for i, header in enumerate(headers):
                    if i < len(row):
                        contact[header.lower().strip()] = row[i].strip()
                contacts.append(contact)
        
        return contacts

    def format_procedure_table(self, table):
        """Format procedure table for display"""
        if not table:
            return []
        
        headers = table[0]
        procedures = []
        
        for row in table[1:]:
            if any(cell.strip() for cell in row):  # Skip empty rows
                procedure = {}
                for i, header in enumerate(headers):
                    if i < len(row):
                        procedure[header.lower().strip()] = row[i].strip()
                procedures.append(procedure)
        
        return procedures

    def get_business_line_from_filename(self, filename):
        """Extract business line identifier from filename"""
        filename_lower = filename.lower()
        
        if 'ce' in filename_lower and 'global' in filename_lower:
            return 'ce'
        elif 'cms' in filename_lower and 'global' in filename_lower:
            return 'cms'
        elif 'cx' in filename_lower and 'tac' in filename_lower:
            return 'cx-tac'
        elif 'cx' in filename_lower and 'labs' in filename_lower:
            return 'cx-labs'
        elif 'lsc' in filename_lower:
            return 'lsc'
        elif 'proactive' in filename_lower:
            return 'proactive'
        elif 'sourced' in filename_lower:
            return 'sourced'
        elif 'tac' in filename_lower and 'global' in filename_lower:
            return 'tac'
        else:
            return filename.replace('.docx', '').lower().replace(' ', '-')

    def convert_document(self, file_path):
        """Convert a single document to JSON format"""
        print(f"Processing: {file_path.name}")
        
        # Extract content
        content = self.extract_text_from_docx(file_path)
        if not content:
            return None
        
        # Categorize content
        categorized = self.categorize_content(content)
        
        # Get business line
        business_line = self.get_business_line_from_filename(file_path.name)
        
        # Create document metadata
        document_data = {
            'metadata': {
                'filename': file_path.name,
                'business_line': business_line,
                'last_modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'converted_date': datetime.now().isoformat()
            },
            'sections': categorized,
            'raw_content': content  # Keep original structure for reference
        }
        
        return document_data

    def convert_all_documents(self):
        """Convert all Word documents in the input directory"""
        docx_files = list(self.input_dir.glob("*.docx"))
        
        if not docx_files:
            print(f"No .docx files found in {self.input_dir}")
            return
        
        print(f"Found {len(docx_files)} documents to convert")
        
        all_documents = {}
        
        for file_path in docx_files:
            if file_path.name.startswith('~'):  # Skip temporary files
                continue
                
            doc_data = self.convert_document(file_path)
            if doc_data:
                business_line = doc_data['metadata']['business_line']
                all_documents[business_line] = doc_data
        
        # Save individual files
        for business_line, doc_data in all_documents.items():
            output_file = self.output_dir / f"{business_line}.json"
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, indent=2, ensure_ascii=False)
            print(f"Saved: {output_file}")
        
        # Save combined index
        index_data = {
            'last_updated': datetime.now().isoformat(),
            'documents': {
                line: {
                    'title': data['metadata']['filename'].replace('.docx', ''),
                    'filename': data['metadata']['filename'],
                    'last_modified': data['metadata']['last_modified'],
                    'sections_count': {
                        section: len(content) for section, content in data['sections'].items()
                    }
                }
                for line, data in all_documents.items()
            }
        }
        
        index_file = self.output_dir / "index.json"
        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(index_data, f, indent=2, ensure_ascii=False)
        print(f"Saved index: {index_file}")
        
        print(f"\nConversion complete! {len(all_documents)} documents converted.")
        return all_documents

def main():
    parser = argparse.ArgumentParser(description='Convert BCP Word documents to JSON')
    parser.add_argument('input_dir', help='Directory containing .docx files')
    parser.add_argument('--output', '-o', default='data', help='Output directory (default: data)')
    
    args = parser.parse_args()
    
    converter = BCPDocumentConverter(args.input_dir, args.output)
    converter.convert_all_documents()

if __name__ == "__main__":
    main()
